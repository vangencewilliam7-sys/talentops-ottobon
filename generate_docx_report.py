import json
from docx import Document
from docx.shared import Pt
import os

def generate_report(json_path, output_path):
    with open(json_path, 'r') as f:
        data = json.load(f)

    doc = Document()
    doc.add_heading('TalentOps AI Latency Report', 0)

    doc.add_heading('Summary', level=1)
    p = doc.add_paragraph('This report explains how fast the AI responds to your questions. We tested three ways the AI works: SLM (General Chat), LLM (Complex Logic), and RAG (Document Search).')

    doc.add_heading('Component Performance', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Total Time (Avg)'
    hdr_cells[2].text = 'Wait Time (Avg)'
    hdr_cells[3].text = 'Speed (Words/Sec)'

    # Extracting recent metrics
    metrics = data.get('latency_metrics', [])
    slm_metrics = [m for m in metrics if m['model_type'] == 'SLM']
    llm_metrics = [m for m in metrics if m['model_type'] == 'LLM']
    rag_metrics = [m for m in metrics if m['model_type'] == 'RAG']

    def get_avg(met_list, key):
        if not met_list: return 0
        return sum(m[key] for m in met_list) / len(met_list)

    components = [
        ('SLM (General Chat)', slm_metrics),
        ('LLM (Complex Logic)', llm_metrics),
        ('RAG (Document Search)', rag_metrics)
    ]

    for name, m_list in components:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = f"{get_avg(m_list, 'total_latency'):.2f}s"
        row_cells[2].text = f"{get_avg(m_list, 'ttft'):.2f}s"
        row_cells[3].text = f"{get_avg(m_list, 'tokens_per_second'):.1f}"

    doc.add_heading('What does this mean?', level=1)
    doc.add_paragraph('• General Chat (SLM): This is our fastest mode. It starts replying almost instantly and writes very quickly.')
    doc.add_paragraph('• Complex Logic (LLM): This takes a bit longer to start thinking but provides detailed answers.')
    doc.add_paragraph('• Document Search (RAG): This spends about 0.6 seconds finding the right documents before it starts writing.')

    doc.add_heading('Bottlenecks (Simple)', level=1)
    doc.add_paragraph('Currently, the system is like a single-lane road. It works perfectly for up to 10 people at once. If 50 people try to use it simultaneously, it gets very slow and starts falling behind. We recommend keeping the number of simultaneous users low or adding more "lanes" (scaling the server).')

    doc.save(output_path)
    print(f"Report saved to {output_path}")

if __name__ == "__main__":
    generate_report('audit_results.json', 'TalentOps_Latency_Report.docx')
