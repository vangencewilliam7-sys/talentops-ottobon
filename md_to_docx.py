#!/usr/bin/env python3
"""
Markdown to DOCX Converter for TalentOps Documentation
Converts markdown files to Microsoft Word .docx format with proper formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import sys
import os

def convert_markdown_to_docx(md_file_path, output_docx_path=None):
    """
    Convert markdown file to formatted DOCX
    
    Args:
        md_file_path: Path to input .md file
        output_docx_path: Path to output .docx file (optional, defaults to same name)
    """
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process markdown line by line
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_language = None
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                # Start of code block
                in_code_block = True
                code_language = line[3:].strip()
                code_lines = []
            else:
                # End of code block
                in_code_block = False
                # Add code block to document
                code_para = doc.add_paragraph('\n'.join(code_lines))
                code_para.style = 'Normal'
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            
            if level == 1:
                heading = doc.add_heading(heading_text, 0)
            else:
                heading = doc.add_heading(heading_text, level)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '___', '***']:
            doc.add_paragraph('_' * 80)
            i += 1
            continue
        
        # Handle lists
        if line.strip().startswith(('- ', '* ', '+ ')):
            list_text = line.strip()[2:]
            # Remove markdown bold/italic
            list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)
            list_text = re.sub(r'\*(.*?)\*', r'\1', list_text)
            doc.add_paragraph(list_text, style='List Bullet')
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            list_text = re.sub(r'^\d+\.\s', '', line.strip())
            list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)
            list_text = re.sub(r'\*(.*?)\*', r'\1', list_text)
            doc.add_paragraph(list_text, style='List Number')
            i += 1
            continue
        
        # Handle tables
        if '|' in line and i > 0 and '|' in lines[i-1]:
            # This is part of a table - collect all table lines
            table_lines = []
            start_idx = i - 1
            while start_idx >= 0 and '|' in lines[start_idx]:
                start_idx -= 1
            start_idx += 1
            
            end_idx = i
            while end_idx < len(lines) and '|' in lines[end_idx]:
                end_idx += 1
            
            table_lines = lines[start_idx:end_idx]
            
            # Parse table
            rows = []
            for tline in table_lines:
                if re.match(r'^\|[\s\-:]+\|', tline):  # Skip separator line
                    continue
                cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                rows.append(cells)
            
            if rows:
                # Create table
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = cell_data
            
            # Skip to end of table
            i = end_idx
            continue
        
        # Handle regular paragraphs
        if line.strip():
            # Remove markdown formatting
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
            line = re.sub(r'\*(.*?)\*', r'\1', line)      # Italic
            line = re.sub(r'`(.*?)`', r'\1', line)        # Inline code
            line = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1', line)  # Links
            
            para = doc.add_paragraph(line)
        
        i += 1
    
    # Determine output path
    if output_docx_path is None:
        base_name = os.path.splitext(md_file_path)[0]
        output_docx_path = f"{base_name}.docx"
    
    # Save document
    doc.save(output_docx_path)
    print(f"✅ Converted: {md_file_path} → {output_docx_path}")
    return output_docx_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(input_file):
        print(f"❌ Error: File not found: {input_file}")
        sys.exit(1)
    
    convert_markdown_to_docx(input_file, output_file)
