#!/usr/bin/env python3
"""
Convert Markdown to Word Document with formatting preservation
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_border_to_table(table):
    """Add borders to table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def parse_markdown_table(lines):
    """Parse markdown table into rows"""
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|---'):
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for c in cells if c]  # Remove empty cells
            if cells:
                rows.append(cells)
    return rows

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    table_buffer = []
    in_code_block = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F5F5F5')
                    run._element.get_or_add_rPr().append(shading_elm)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and (i == 0 or '|' in lines[i-1] or (i < len(lines)-1 and '|' in lines[i+1])):
            table_buffer.append(line)
            i += 1
            # Check if next line is also part of table
            if i < len(lines) and '|' not in lines[i]:
                # Process table
                table_rows = parse_markdown_table(table_buffer)
                if table_rows:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            
                            # Bold header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    
                    add_border_to_table(table)
                    doc.add_paragraph()  # Add spacing after table
                
                table_buffer = []
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle blockquotes
        elif line.startswith('>'):
            quote_text = line[1:].strip()
            p = doc.add_paragraph(quote_text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(96, 96, 96)
        
        # Handle lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Check for checkboxes
            if text.startswith('✅'):
                text = '✓ ' + text[2:]
            elif text.startswith('❌'):
                text = '✗ ' + text[2:]
            doc.add_paragraph(text, style='List Bullet')
        
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            doc.add_paragraph(text, style='List Number')
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        
        # Handle bold/italic inline formatting
        elif line.strip():
            # Process inline formatting
            text = line
            text = re.sub(r'\*\*\*(.+?)\*\*\*', r'***\1***', text)  # Bold+Italic
            text = re.sub(r'\*\*(.+?)\*\*', r'**\1**', text)  # Bold
            text = re.sub(r'\*(.+?)\*', r'*\1*', text)  # Italic
            text = re.sub(r'`(.+?)`', r'`\1`', text)  # Code
            
            p = doc.add_paragraph()
            
            # Split and format
            parts = re.split(r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)', text)
            for part in parts:
                if part.startswith('***') and part.endswith('***'):
                    run = p.add_run(part[3:-3])
                    run.bold = True
                    run.italic = True
                elif part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    p.add_run(part)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    md_file = "/Users/kummaraakash/.gemini/antigravity/brain/b69d690e-949c-4b29-bb0d-43d056af1427/chatbot_implementation_guide.md.resolved"
    docx_file = "/Users/kummaraakash/Desktop/TalentOps_Chatbot_Implementation_Guide.docx"
    
    convert_md_to_docx(md_file, docx_file)
