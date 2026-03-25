import io
import logging
import httpx
import asyncio
import PyPDF2
from typing import List, Optional

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

logger = logging.getLogger(__name__)

# Global client to be initialized by the host (should be an AsyncOpenAI client)
openai_client = None

def init_rag(client):
    """Initialize RAG utilities with an AsyncOpenAI client."""
    global openai_client
    openai_client = client
    logger.info("RAG utilities initialized with AsyncOpenAI client.")

import re

def chunk_text(text: str, max_chunk_words=300, overlap_words=75) -> List[str]:
    """Semantic chunking: Split text prioritizing paragraphs and sentences."""
    if not text.strip():
        return []
        
    # Standardize whitespace but keep paragraphs
    text = re.sub(r'\r\n', '\n', text)
    paragraphs = re.split(r'\n{2,}', text)
    
    chunks = []
    current_chunk = ""
    current_word_count = 0
    
    for para in paragraphs:
        para = para.strip()
        if not para: continue
        
        # Split paragraph into sentences roughly
        sentences = re.split(r'(?<=[.!?])\s+', para)
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence: continue
            
            words = sentence.split()
            word_count = len(words)
            
            # If a single sentence is huge, we have to split it naively
            if word_count > max_chunk_words:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                    current_word_count = 0
                
                # Naive split for the giant sentence
                for i in range(0, word_count, max_chunk_words - overlap_words):
                    sub_chunk = " ".join(words[i:i + max_chunk_words])
                    chunks.append(sub_chunk)
                continue
            
            # If adding this sentence exceeds the limit, save current chunk and start a new one
            if current_word_count + word_count > max_chunk_words and current_word_count > 0:
                chunks.append(current_chunk.strip())
                
                # Calculate overlap: take the last `overlap_words` from the current chunk
                current_words = current_chunk.split()
                overlap_text = " ".join(current_words[-overlap_words:]) if len(current_words) > overlap_words else current_chunk
                
                current_chunk = overlap_text + " " + sentence
                current_word_count = len(current_chunk.split())
            else:
                current_chunk = (current_chunk + " " + sentence).strip()
                current_word_count += word_count
                
        # Add paragraph break if the chunk isn't empty
        if current_chunk:
            current_chunk += "\n\n"

    # Don't forget the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
        
    return chunks

async def get_embeddings(texts: List[str]) -> List[List[float]]:
    """Generate embeddings for a list of texts using the initialized AsyncOpenAI client."""
    if openai_client is None:
        logger.error("RAG utilities not initialized. Call init_rag(client) first.")
        return [], 0.0
        
    clean_texts = [t.replace("\n", " ") for t in texts]
    if not clean_texts:
        return [], 0.0
    
    import time
    start_time = time.perf_counter()
    try:
        resp = await openai_client.embeddings.create(
            input=clean_texts, 
            model="text-embedding-3-small",
            timeout=10
        )
        latency = time.perf_counter() - start_time
        return [d.embedding for d in resp.data], latency
    except Exception as e:
        logger.error(f"Embedding Error: {e}")
        return [], 0.0

async def parse_file_from_url(url: str) -> str:
    """Download and parse a file (PDF, DOCX, or text) from a URL asynchronously."""
    if not url: return ""
    try:
        logger.info(f"Downloading for RAG (Async): {url}")
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.get(url)
            resp.raise_for_status()
            
            content_type = resp.headers.get('Content-Type', '').lower()
            url_lower = url.lower()
            file_stream = io.BytesIO(resp.content)

            if url_lower.endswith('.pdf') or 'pdf' in content_type:
                reader = PyPDF2.PdfReader(file_stream)
                text = ""
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
                return text
            
            elif url_lower.endswith('.docx') or 'wordprocessingml' in content_type:
                if DocxDocument is None:
                    logger.warning("python-docx not installed. Skipping DOCX.")
                    return ""
                
                # --- EXHAUSTIVE EXTRACTION ---
                # Basic docx.Document() misses text in shapes, textboxes, and some nested structures.
                # We use a direct XML 't' (text) element scrape for maximum coverage.
                doc = DocxDocument(file_stream)
                text_parts = []
                
                # 1. Standard Paragraphs
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text_parts.extend(paragraphs)
                
                # 2. Standard Tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text: text_parts.append(row_text)
                
                # 3. Deep XML Scrape (Safety Fallback for Textboxes/Shapes)
                try:
                    all_xml_text = []
                    # Find ALL 'w:t' text elements in the entire body XML
                    for t in doc.element.xpath('//w:t'):
                        if t.text: all_xml_text.append(t.text.strip())
                    
                    combined_xml = " ".join(all_xml_text)
                    combined_parts = " ".join(text_parts)
                    
                    # If XML scrape found significantly more text, it means we missed things (shapes/boxes)
                    if len(combined_xml) > len(combined_parts) * 1.2:
                        logger.info(f"⚡ DOCX: Deep XML scrape found extra content ({len(combined_xml)} vs {len(combined_parts)} chars)")
                        return "\n\n".join(text_parts) + "\n\n--- ADDITIONAL CONTENT ---\n\n" + combined_xml
                except Exception as xml_e:
                    logger.warning(f"XML scrape failed: {xml_e}")

                return "\n\n".join(text_parts)
                
            else:
                try:
                    return resp.content.decode('utf-8')
                except:
                    return resp.content.decode('latin-1', errors='ignore')
    except Exception as e:
        logger.error(f"File Parse Error: {e}")
        return ""
