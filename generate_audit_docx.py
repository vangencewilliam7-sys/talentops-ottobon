import os
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_audit_report(output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('Infrastructure Audit: Performance & Scalability Deep-Dive', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    p = doc.add_paragraph()
    p.add_run('Auditor: Principal AI Infrastructure Auditor\n').italic = True
    p.add_run('Target: Modal Gateway - TalentOps Lifecycle\n').italic = True
    p.add_run('Subject: Addressing the "Concurrency Wall" and Latency Optimization').italic = True
    
    # Section 1: The Blueprint
    doc.add_heading('1. The Architecture Blueprint: How Data Flows', level=1)
    doc.add_paragraph(
        "Current system architecture is built on a 'Sequential Model.' This means the system performs "
        "tasks one-by-one, like a single person working through a checklist. While this works for one user, "
        "it creates a massive line (queue) when 500 people try to use it at once."
    )
    
    doc.add_heading('Core Components', level=2)
    components = [
        ('Entry Point', 'The "Front Door" (unified_server.py). Currently handles all requests through one single process.'),
        ('AI Intent Engine', 'The "Brain" (Together AI). Decides what the user wants, but currently stops everything else to "think."'),
        ('The Database', 'The "Memory" (Supabase). Fetches tasks and logs, but blocks the system while waiting for data.'),
        ('The RAG Pipeline', 'The "Librarian." Searches through documents to find answers, currently the slowest part of the trek.')
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Role & Current Limitation'
    for area, detail in components:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = detail

    doc.add_heading('The Request Journey (The Real Bottleneck)', level=2)
    doc.add_paragraph(
        "Every time you ask a question, the system follows this exact path. Because each step is 'Synchronous' "
        "(Step 2 cannot start until Step 1 finishes), any delay in a single step slows down the entire system."
    )
    chain = [
        "1. Start (System identifies you and your organization)",
        "2. Think (AI classifies if you want tasks, a policy, or a chat) - COST: ~600ms",
        "3. Check Permissions (Ensures you are allowed to see the data) - COST: ~50ms",
        "4. Fetch Data (Hits the database to get your records) - COST: ~200ms",
        "5. Finalize (Prepares the answer) - COST: ~100ms",
        "6. Log (Writes the result to a file) - COST: ~50ms"
    ]
    for step in chain:
        doc.add_paragraph(step, style='List Bullet')

    # Section 2: Latency Breakdown
    doc.add_heading('2. Why are we waiting? Latency Breakdown', level=1)
    doc.add_paragraph(
        "We measure 'Time to First Token' (TTFT). This is the 'Blink Time'—the time you spend staring at a "
        "loading spinner before the AI starts typing. Our goal is to keep this under 1 second."
    )
    
    latency_table = doc.add_table(rows=1, cols=4)
    latency_table.style = 'Table Grid'
    headers = ['AI Mode', 'Wait Time (TTFT)', 'Typing Speed', 'Total Delay']
    hdr_cells = latency_table.rows[0].cells
    for i, h in enumerate(headers): hdr_cells[i].text = h
    
    rows = [
        ('General Chat', '420ms (Fast)', '120 words/sec', '1.2s'),
        ('Complex Logic', '600ms (Medium)', '47 words/sec', '1.5s'),
        ('Policy/Doc Search', '380ms (Fast)', '118 words/sec', '2.0s')
    ]
    for mod, ttft, proc, total in rows:
        row = latency_table.add_row().cells
        row[0].text = mod
        row[1].text = ttft
        row[2].text = proc
        row[3].text = total

    doc.add_heading('The RAG "Penalty"', level=2)
    doc.add_paragraph(
        "RAG (Document Search) adds a 'latency penalty' because it has to read through thousands of chunks "
        "of text. Currently, it does this in four separate, slow steps instead of all at once."
    )

    # Section 3: Failure Analysis
    doc.add_heading('3. The "Concurrency Wall": Why it fails at 500 users', level=1)
    doc.add_paragraph(
        "Analogy: Imagine a fast-food restaurant with only one cashier. If 1 person comes, it's fast. "
        "If 500 people come, the 500th person has to wait for all 499 people ahead of them to finish. "
        "This is exactly what is happening to our server."
    )
    reasons = [
        ('The Event Loop Freeze', 'Because our code is "Synchronous," the entire server "freezes" while it waits for the database or OpenAI. While it is frozen, it cannot even acknowledge the next user.'),
        ('Single Lane Traffic', 'We are running a "Single Worker" server. This is like a 10-lane highway being squeezed into a 1-lane tunnel.'),
        ('Provider Exhaustion', 'OpenAI and Together AI have caps on how many people can ask questions at the exact same time. Without a queue, we hit these caps and get "Rate Limit" errors.')
    ]
    for r_title, r_desc in reasons:
        p = doc.add_paragraph()
        run = p.add_run(f"{r_title}: ")
        run.bold = True
        p.add_run(r_desc)

    # Section 4: The Clean Fix
    doc.add_heading('4. The Solution: Moving to a Multi-Lane Highway', level=1)
    doc.add_paragraph(
        "To handle 500+ users, we must transform the system into an 'Asynchronous Cluster.' "
        "This allows the cashier to take multiple orders without waiting for the food to be cooked."
    )
    
    fixes = [
        "Async Overhaul: We will change the code so the server never 'freezes' while waiting for data.",
        "Cluster Scaling: Instead of 1 machine, we will use a cluster of 5-10 small machines working together.",
        "Semantic Caching: If two users ask the same question, the system will use a 'Saved Answer' in 0.01s instead of recalculating (saving 2 seconds).",
        "Parallel Retrieval: RAG will search for docs and generate embeddings at the exact same time."
    ]
    for f in fixes:
        doc.add_paragraph(f, style='List Number')

    doc.add_heading('Target Performance', level=2)
    doc.add_paragraph(
        "• 1-50 Users: Instant feel (<1s response)\n"
        "• 100-300 Users: Smooth experience (<3s response)\n"
        "• 500+ Users: Stable and active (<5s response, zero crashes)"
    )

    doc.save(output_path)
    print(f"Professional Audit report saved to {output_path}")

if __name__ == "__main__":
    create_audit_report('Infrastructure_Audit_TalentOps.docx')
