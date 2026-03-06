import os
import re
from docx import Document
from docx.shared import Pt

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip('\n')
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            # Simple bold/italic parsing
            add_formatted_text(p, text)
        
        # Paragraphs
        elif line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        
        # Empty lines
        else:
            pass

    doc.save(docx_path)
    print(f"Converted {md_path} to {docx_path}")

def add_formatted_text(paragraph, text):
    # Very simple regex for bold **text**
    # Splitting by bold tags while keeping the tags
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    base_path = r"C:\Users\DELL\.gemini\antigravity\brain\12537d4b-fb6c-4550-84cf-4d0946d1b158"
    
    # Dynamically find all .md files (excluding task and implementation_plan)
    files = [f for f in os.listdir(base_path) if f.endswith('.md') and f not in ["task.md", "implementation_plan.md"]]
    
    for f in files:
        md = os.path.join(base_path, f)
        docx = os.path.join(base_path, f.replace(".md", ".docx"))
        md_to_docx(md, docx)
